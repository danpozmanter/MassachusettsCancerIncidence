from collections import OrderedDict
import json
from docx import Document
import config
from models import DataJSONEncoder, CancerStats, CancerData


"""General note:
Certain aspects of these documents remain stable
(paragraph style for city name, location of columns and rows within tables).

If these documents are updated in the future this might change.
"""


def get_document(docx_path: str) -> Document:
    """Read the word doc, returning a parsed Document instance."""
    with open(docx_path, "rb") as f:
        return Document(f)


def get_cities(doc: Document) -> list[str]:
    """Get the list of cities in the word doc.
    In these documents, Heading 1 is only used to denote city titles."""
    headings = []
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            headings.append(p.text)
    return headings


def get_cancer_stats(row, offset):
    """Pull cancer stats for a particular row.
    Parse out "nc" (not calculated), and represent this as null.
    """
    if len(row.cells[offset].text.strip()) > 0:
        sir = row.cells[offset+2].text
        ci = row.cells[offset+3].text
        if ci and "nc" not in ci:
            ci = ci.strip("()").split("-")
            ci_low = float(ci[0])
            ci_high = float(ci[1])
            significant = True
        else:
            ci_low = None
            ci_high = None
            significant = False
        return CancerStats(
            observed=float(row.cells[offset].text),
            expected=float(row.cells[offset+1].text),
            sir=float(sir) if sir != "nc" else None,
            ci_low=ci_low,
            ci_high=ci_high,
            significant=significant
        )
    return CancerStats(0, 0, None, None, None, False)


def get_cancer_data(male_row, female_row, offset=0) -> CancerData:
    """Pull cancer data for a category of cancer.
    Row data is divded by sex.
    Create "combined" data, leaving out confidence intervals.
    Include a SIR if both the male and female data was significant.
    (Otherwise that would skew the data)."""
    male_stats = get_cancer_stats(male_row, offset)
    female_stats = get_cancer_stats(female_row, offset)
    combined_observed = male_stats.observed + female_stats.observed
    combined_expected = male_stats.expected + female_stats.expected
    combinded_significant = (
        male_stats.significant and female_stats.significant)
    if combinded_significant:
        combined_sir = (combined_observed / combined_expected) * 100
    else:
        combined_sir = None
    combined_stats = CancerStats(
        observed=combined_observed,
        expected=combined_expected,
        sir=combined_sir,
        ci_low=None,
        ci_high=None,
        significant=combinded_significant
    )
    return CancerData(
        male=male_stats,
        female=female_stats,
        combined=combined_stats
    )


def process_city(year_range: str, city: str, table):
    """Cancer data is presented in tables split on left and right.
    Iterate through the rows
    """
    print(f"    Processing... {city}")
    cancers = OrderedDict()
    for r in range(1, len(table.rows), 3):
        cancer_left = table.rows[r].cells[0].text
        cancer_right = table.rows[r].cells[6].text
        cancers[cancer_left] = get_cancer_data(
            table.rows[r+1],
            table.rows[r+2],
            1
        )
        cancers[cancer_right] = get_cancer_data(
            table.rows[r+1],
            table.rows[r+2],
            7
        )
    with open(f"city_json/{year_range}/{city}.json", "w") as f:
        f.write(json.dumps(cancers, cls=DataJSONEncoder))


def process_doc(year_range: str, docx_path: str):
    """Process a word doc containing cancer data broken down by city.
    Parse out the tables and associated city names, and process by city."""
    doc = get_document(docx_path)
    tables = doc.tables
    cities = get_cities(doc)
    if len(tables) != len(cities):
        raise Exception(f"For {docx_path} cities and tables do not match")
    for i, city in enumerate(cities):
        process_city(year_range, city, tables[i])


for year_range in config.year_ranges:
    print(f"Processing year range: {year_range}")
    for city_group in config.city_groups:
        print(f"Processing city group: {city_group}")
        process_doc(
            year_range,
            f"raw/{year_range}/registry-city-{year_range}-{city_group}.docx"
        )
